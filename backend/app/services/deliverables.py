"""보고서 산출물 생성 — 사진대지(Word) · 손상물량표(Excel) · 외관조사망도(DXF).

KO-Detect Compact QuickGuide STEP 05가 규정한 세 가지 결과 파일을 만든다.
사용자는 최소 1개를 골라 받고, 여러 개면 ZIP 하나로 묶인다.

| 결과 파일 | 형식 | 내용 |
|---|---|---|
| 사진대지 | Word | 진단대상 사진을 문서로 정리, 손상 상세 이미지 목록 포함 |
| 손상물량표 | Excel | 손상 종류와 수량을 표로 정리 |
| 외관조사망도 | DXF | 도면에 위치·사진·손상물량표를 그려 산출 |

설계 원칙
---------
* **물량은 실측 단위로 낸다.** 균열은 연장(m), 면적형 결함은 면적(m²), 개소형은
  개소(EA). 세 가지를 한 칸에 섞으면 집계가 불가능해진다.
* **AI 결과와 사람이 그린 것을 구분해 표기한다.** 검토자가 무엇을 다시 봐야
  하는지 알아야 한다.
* **스케일이 없으면 물량을 내지 않는다.** 픽셀을 m로 환산할 수 없으면 "산출 불가"로
  적는다. 임의 가정값으로 채우면 그 숫자가 보고서에 그대로 실린다.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from ..domain import DEFECT_LABELS_KO, MEMBER_CLASSES, DefectType

# ─── 물량 단위 ─────────────────────────────────────────────────
# 결함 유형별 산출 단위. 균열만 연장(m)이고 나머지는 면적 또는 개소다.
QUANTITY_UNIT: dict[DefectType, str] = {
    DefectType.CRACK: "m",
    DefectType.SPALLING: "m2",
    DefectType.EFFLORESCENCE: "m2",
    DefectType.LEAKAGE: "m2",
    DefectType.REBAR_EXPOSURE: "EA",
    DefectType.SEGREGATION: "m2",
    DefectType.DAMAGE: "EA",
}

UNIT_LABEL_KO = {"m": "연장 (m)", "m2": "면적 (m²)", "EA": "개소 (EA)"}


@dataclass
class DefectRow:
    """산출물에 들어가는 결함 1건."""

    defect_id: str
    defect_type: DefectType
    member_code: str
    group_name: str
    spot_label: str
    grade: str
    width_mm: float | None
    length_mm: float | None
    area_ratio: float | None
    photo_name: str
    overlay_path: Path | None
    source: str = "ai"            # ai | manual | edited
    confidence: float = 1.0
    basis: str = ""
    # 사진의 실치수 환산 계수. 없으면 물량 산출 불가.
    mm_per_px: float | None = None
    photo_area_px: int = 0

    @property
    def unit(self) -> str:
        return QUANTITY_UNIT.get(self.defect_type, "EA")

    @property
    def photo_area_m2(self) -> float | None:
        """사진 한 장이 담은 실제 벽면 면적."""
        if self.mm_per_px is None or not self.photo_area_px:
            return None
        return self.photo_area_px * (self.mm_per_px / 1000.0) ** 2

    @property
    def quantity(self) -> float | None:
        """단위에 맞는 물량. 스케일이 없으면 None(산출 불가)."""
        if self.unit == "EA":
            return 1.0
        if self.unit == "m":
            return None if self.length_mm is None else round(self.length_mm / 1000.0, 3)
        # 면적형 — 사진 면적률 × 사진이 담은 실제 면적
        if self.area_ratio is None:
            return None
        area = self.photo_area_m2
        return None if area is None else round(self.area_ratio * area, 4)

    @property
    def quantity_basis(self) -> str:
        """물량 산출 근거 — 검토자가 숫자를 되짚을 수 있어야 한다.

        면적형 물량은 '사진이 담은 면적'에 좌우된다. 접사(0.2mm/px)와 원거리
        촬영(6mm/px)은 같은 면적률이라도 물량이 700배 차이 난다. 근거를 적지
        않으면 실무자가 이 차이를 알아채지 못한다.
        """
        if self.unit == "EA":
            return "관측 1개소"
        if self.unit == "m":
            if self.length_mm is None:
                return "길이 미측정 — 산출 불가"
            return f"검출 연장 {self.length_mm:.0f}mm"
        if self.area_ratio is None:
            return "면적률 미산출 — 산출 불가"
        area = self.photo_area_m2
        if area is None:
            return "스케일 미확정 — 산출 불가"
        return (
            f"면적률 {self.area_ratio * 100:.2f}% × 사진면적 {area:.3f}㎡ "
            f"(GSD {self.mm_per_px:.3f}mm/px)"
        )


@dataclass
class ReportBundle:
    """산출물 생성에 필요한 입력 일체."""

    building_name: str
    facility_class: str
    inspected_at: datetime
    inspector: str
    safety_grade: str | None
    rows: list[DefectRow]
    groups: list[dict] = field(default_factory=list)   # {name, member_code, photos}
    drawings: list[dict] = field(default_factory=list) # {name, width, height, spots}
    scope_label: str = "시설물·도면 기준"


# ─── 손상물량표 (Excel) ────────────────────────────────────────
def build_quantity_xlsx(b: ReportBundle) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="14181F")
    head_font = Font(color="FFFFFF", bold=True, size=10)
    title_font = Font(bold=True, size=14)

    # ── 시트 1: 집계 ──
    ws = wb.active
    ws.title = "집계"
    ws["A1"] = "손상물량표"
    ws["A1"].font = title_font
    ws["A2"] = (
        f"{b.building_name} · {b.facility_class} · "
        f"점검일 {b.inspected_at:%Y-%m-%d} · 점검자 {b.inspector or '—'}"
    )
    ws["A3"] = (
        f"종합 안전등급 {b.safety_grade or '—'} · "
        f"발행 {datetime.now():%Y-%m-%d %H:%M}"
    )
    ws["A4"] = "AI 분석 결과는 참고용입니다. 손상 수치와 최종 내용은 반드시 확인하십시오."
    ws["A4"].font = Font(color="C0392B", size=9)
    ws["A5"] = (
        "면적형 물량은 [면적률 × 사진이 담은 실제 면적]으로 산출합니다. "
        "촬영 거리에 따라 사진 면적이 크게 달라지므로 '상세' 시트의 산출근거를 함께 확인하십시오."
    )
    ws["A5"].font = Font(color="5B6472", size=8)

    headers = ["결함유형", "단위", "건수", "물량 합계", "산출 불가", "AI 검출", "직접 입력"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=7, column=c, value=h)
        cell.fill, cell.font, cell.border = head_fill, head_font, border
        cell.alignment = Alignment(horizontal="center")

    by_type: dict[DefectType, list[DefectRow]] = {}
    for r in b.rows:
        by_type.setdefault(r.defect_type, []).append(r)

    row_i = 8
    for dtype in DefectType:
        rows = by_type.get(dtype, [])
        if not rows:
            continue
        unit = QUANTITY_UNIT[dtype]
        qtys = [r.quantity for r in rows]
        known = [q for q in qtys if q is not None]
        vals = [
            DEFECT_LABELS_KO[dtype],
            UNIT_LABEL_KO[unit],
            len(rows),
            round(sum(known), 3) if known else "산출 불가",
            sum(1 for q in qtys if q is None),
            sum(1 for r in rows if r.source == "ai"),
            sum(1 for r in rows if r.source != "ai"),
        ]
        for c, v in enumerate(vals, 1):
            cell = ws.cell(row=row_i, column=c, value=v)
            cell.border = border
            if c >= 3:
                cell.alignment = Alignment(horizontal="right")
        row_i += 1

    total = ws.cell(row=row_i, column=1, value="합계")
    total.font = Font(bold=True)
    total.border = border
    ws.cell(row=row_i, column=3, value=len(b.rows)).font = Font(bold=True)
    for c in range(1, 8):
        ws.cell(row=row_i, column=c).border = border

    for col, w in zip("ABCDEFG", (16, 14, 10, 14, 12, 12, 12)):
        ws.column_dimensions[col].width = w

    # ── 시트 2: 상세 ──
    ws2 = wb.create_sheet("상세")
    detail_headers = [
        "결함코드", "유형", "부재", "사진그룹", "위치", "등급",
        "폭(mm)", "길이(mm)", "물량", "단위", "물량 산출근거",
        "출처", "신뢰도", "사진", "판정근거",
    ]
    for c, h in enumerate(detail_headers, 1):
        cell = ws2.cell(row=1, column=c, value=h)
        cell.fill, cell.font, cell.border = head_fill, head_font, border
        cell.alignment = Alignment(horizontal="center")

    for i, r in enumerate(b.rows, start=2):
        member = MEMBER_CLASSES.get(r.member_code)
        q = r.quantity
        vals = [
            r.defect_id,
            DEFECT_LABELS_KO.get(r.defect_type, r.defect_type.value),
            member.label_ko if member else r.member_code,
            r.group_name or "미분류",
            r.spot_label or "—",
            (r.grade or "").upper(),
            round(r.width_mm, 3) if r.width_mm is not None else "—",
            round(r.length_mm, 1) if r.length_mm is not None else "—",
            q if q is not None else "산출 불가",
            r.unit,
            r.quantity_basis,
            {"ai": "AI", "manual": "직접입력", "edited": "AI+수정"}.get(r.source, r.source),
            round(r.confidence, 3),
            r.photo_name,
            r.basis,
        ]
        for c, v in enumerate(vals, 1):
            cell = ws2.cell(row=i, column=c, value=v)
            cell.border = border
            if c in (7, 8, 9, 13):
                cell.alignment = Alignment(horizontal="right")

    for col, w in zip(
        ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O"],
        (18, 12, 12, 14, 14, 8, 11, 12, 11, 9, 42, 11, 10, 26, 30),
    ):
        ws2.column_dimensions[col].width = w
    ws2.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── 사진대지 (Word) ───────────────────────────────────────────
def build_photo_sheet_docx(b: ReportBundle) -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    title = doc.add_heading("사진대지", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"{b.building_name} · {b.facility_class}\n"
        f"점검일 {b.inspected_at:%Y-%m-%d} · 점검자 {b.inspector or '—'} · "
        f"종합 안전등급 {b.safety_grade or '—'}"
    )
    run.font.size = Pt(10)

    warn = doc.add_paragraph()
    wr = warn.add_run(
        "AI 분석 결과는 참고용입니다. 최종 보고서의 손상 수치와 내용은 "
        "출력 전 반드시 검토하십시오."
    )
    wr.font.size = Pt(9)
    wr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # 그룹 단위로 정리 — QuickGuide: "보고서가 그룹 단위로 정리됩니다"
    by_group: dict[str, list[DefectRow]] = {}
    for r in b.rows:
        by_group.setdefault(r.group_name or "미분류", []).append(r)

    for gname, rows in by_group.items():
        doc.add_page_break()
        doc.add_heading(f"사진그룹 · {gname}", level=1)

        p = doc.add_paragraph()
        pr = p.add_run(f"결함 {len(rows)}건")
        pr.font.size = Pt(9)

        # 2열 그리드로 사진과 캡션을 배치한다
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i in range(0, len(rows), 2):
            img_cells = table.add_row().cells
            cap_cells = table.add_row().cells
            for j in range(2):
                if i + j >= len(rows):
                    continue
                r = rows[i + j]
                cell = img_cells[j]
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if r.overlay_path and r.overlay_path.exists():
                    try:
                        para.add_run().add_picture(str(r.overlay_path), width=Cm(7.4))
                    except Exception:
                        para.add_run("(이미지 삽입 실패)").font.size = Pt(8)
                else:
                    para.add_run("(이미지 없음)").font.size = Pt(8)

                member = MEMBER_CLASSES.get(r.member_code)
                q = r.quantity
                caption = (
                    f"{r.defect_id}  "
                    f"{DEFECT_LABELS_KO.get(r.defect_type, r.defect_type.value)} · "
                    f"{member.label_ko if member else r.member_code}\n"
                    f"등급 {(r.grade or '—').upper()} · "
                    f"폭 {f'{r.width_mm:.2f}mm' if r.width_mm is not None else '—'} · "
                    f"물량 {f'{q} {r.unit}' if q is not None else '산출 불가'}\n"
                    f"위치 {r.spot_label or '—'} · "
                    f"출처 {'AI' if r.source == 'ai' else '직접입력' if r.source == 'manual' else 'AI+수정'}"
                )
                cp = cap_cells[j].paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(caption)
                cr.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 외관조사망도 (DXF) ────────────────────────────────────────
# 결함 유형별 레이어 색상 (AutoCAD Color Index)
DXF_LAYER_COLOR: dict[str, int] = {
    "crack": 1,            # 빨강
    "spalling": 30,        # 주황
    "efflorescence": 7,    # 흰색
    "leakage": 5,          # 파랑
    "rebar_exposure": 6,   # 자홍
    "segregation": 3,      # 초록
    "damage": 2,           # 노랑
}


def build_survey_dxf(b: ReportBundle) -> bytes:
    """도면에 위치 핀·그룹·손상물량표를 그린 DXF를 만든다.

    도면이 없으면 빈 도면틀에 위치를 배치한다 — QuickGuide가 '빈 화면으로
    시작'을 허용하므로 산출물도 그 경우를 지원해야 한다.
    """
    import ezdxf

    doc = ezdxf.new("R2010", setup=True)
    msp = doc.modelspace()

    # 레이어 구성 — 결함 유형별로 분리해 CAD에서 켜고 끌 수 있게 한다
    doc.layers.add("DRAWING_FRAME", color=8)
    doc.layers.add("SPOT_PIN", color=4)
    doc.layers.add("SPOT_TEXT", color=7)
    doc.layers.add("QUANTITY_TABLE", color=7)
    for code, color in DXF_LAYER_COLOR.items():
        doc.layers.add(f"DEFECT_{code.upper()}", color=color)

    y_cursor = 0.0
    for dw in b.drawings or [{"name": "빈 도면", "width": 1600, "height": 1200, "spots": []}]:
        w = float(dw.get("width", 1600))
        h = float(dw.get("height", 1200))
        oy = y_cursor

        # 도면틀
        msp.add_lwpolyline(
            [(0, oy), (w, oy), (w, oy + h), (0, oy + h), (0, oy)],
            dxfattribs={"layer": "DRAWING_FRAME"},
        )
        msp.add_text(
            dw.get("name", "도면"),
            dxfattribs={"layer": "DRAWING_FRAME", "height": h * 0.025},
        ).set_placement((10, oy + h + h * 0.01))

        for sp in dw.get("spots", []):
            x = float(sp.get("x", 0))
            # DXF는 Y축이 위로 증가하므로 화면 좌표를 뒤집는다
            y = oy + h - float(sp.get("y", 0))
            r = h * 0.014

            msp.add_circle((x, y), r, dxfattribs={"layer": "SPOT_PIN"})
            msp.add_text(
                str(sp.get("number", "")),
                dxfattribs={"layer": "SPOT_PIN", "height": r * 1.1},
            ).set_placement((x, y), align=ezdxf.enums.TextEntityAlignment.MIDDLE_CENTER)

            label = (
                f"{sp.get('group_name', '')} "
                f"({sp.get('photo_count', 0)}장, 손상 {sp.get('defect_count', 0)}건)"
            )
            msp.add_text(
                label, dxfattribs={"layer": "SPOT_TEXT", "height": r * 0.85}
            ).set_placement((x + r * 1.4, y - r * 0.4))

            note = " · ".join(
                x for x in (sp.get("member_label", ""), sp.get("direction", "")) if x
            )
            if note:
                msp.add_text(
                    note, dxfattribs={"layer": "SPOT_TEXT", "height": r * 0.7}
                ).set_placement((x + r * 1.4, y - r * 1.5))

        y_cursor += h * 1.15

    # 손상물량표 — 도면 오른쪽에 그린다
    by_type: dict[DefectType, list[DefectRow]] = {}
    for r in b.rows:
        by_type.setdefault(r.defect_type, []).append(r)

    tx = (b.drawings[0]["width"] if b.drawings else 1600) * 1.08
    ty = y_cursor
    line_h = 34.0
    col_w = (200.0, 90.0, 110.0)

    msp.add_text(
        "손상물량표", dxfattribs={"layer": "QUANTITY_TABLE", "height": line_h * 0.8}
    ).set_placement((tx, ty))
    ty -= line_h * 1.4

    for hi, htxt in enumerate(("결함유형", "건수", "물량")):
        msp.add_text(
            htxt, dxfattribs={"layer": "QUANTITY_TABLE", "height": line_h * 0.55}
        ).set_placement((tx + sum(col_w[:hi]), ty))
    ty -= line_h

    for dtype, rows in by_type.items():
        known = [q for q in (r.quantity for r in rows) if q is not None]
        qty_txt = (
            f"{sum(known):.2f} {QUANTITY_UNIT[dtype]}" if known else "산출 불가"
        )
        for ci, txt in enumerate(
            (DEFECT_LABELS_KO[dtype], f"{len(rows)}", qty_txt)
        ):
            msp.add_text(
                txt,
                dxfattribs={
                    "layer": f"DEFECT_{dtype.value.upper()}",
                    "height": line_h * 0.5,
                },
            ).set_placement((tx + sum(col_w[:ci]), ty))
        ty -= line_h

    ty -= line_h * 0.5
    msp.add_text(
        "AI 분석 결과는 참고용 — 최종 수치는 검토 후 확정",
        dxfattribs={"layer": "QUANTITY_TABLE", "height": line_h * 0.45},
    ).set_placement((tx, ty))

    stream = io.StringIO()
    doc.write(stream)
    return stream.getvalue().encode("utf-8")


# ─── ZIP 묶기 ──────────────────────────────────────────────────
def build_bundle_zip(b: ReportBundle, kinds: list[str]) -> tuple[bytes, list[str]]:
    """선택한 결과 파일을 ZIP 하나로 묶는다.

    QuickGuide STEP 05 — "고른 파일이 ZIP 하나로 묶여 저장되고, 끝나면 저장
    폴더가 자동으로 열립니다."
    """
    stamp = b.inspected_at.strftime("%Y%m%d")
    safe = "".join(c for c in b.building_name if c.isalnum() or c in " _-").strip()
    made: list[str] = []

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        if "photo_sheet" in kinds:
            name = f"사진대지_{safe}_{stamp}.docx"
            zf.writestr(name, build_photo_sheet_docx(b))
            made.append(name)
        if "quantity" in kinds:
            name = f"손상물량표_{safe}_{stamp}.xlsx"
            zf.writestr(name, build_quantity_xlsx(b))
            made.append(name)
        if "survey_dxf" in kinds:
            name = f"외관조사망도_{safe}_{stamp}.dxf"
            zf.writestr(name, build_survey_dxf(b))
            made.append(name)

        # 어떤 조건에서 생성됐는지 남긴다 — 나중에 이 ZIP만 보고도 알 수 있게
        manifest = [
            "KO-Detect 보고서 산출물",
            f"시설물     {b.building_name} ({b.facility_class})",
            f"점검일     {b.inspected_at:%Y-%m-%d}",
            f"점검자     {b.inspector or '—'}",
            f"안전등급   {b.safety_grade or '—'}",
            f"보고서범위 {b.scope_label}",
            f"결함 건수  {len(b.rows)}건 "
            f"(AI {sum(1 for r in b.rows if r.source == 'ai')} · "
            f"직접입력 {sum(1 for r in b.rows if r.source != 'ai')})",
            f"생성 시각  {datetime.now():%Y-%m-%d %H:%M:%S}",
            "",
            "포함 파일",
            *[f"  - {n}" for n in made],
            "",
            "AI 분석 결과는 참고용입니다. 최종 보고서의 손상 수치와 내용은",
            "출력 전 반드시 검토하십시오.",
        ]
        zf.writestr("README.txt", "\n".join(manifest).encode("utf-8"))

    return buf.getvalue(), made
